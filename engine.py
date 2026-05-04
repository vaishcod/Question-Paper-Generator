from docx import Document
import PyPDF2
from llm_client import generate_text
import re

# =========================
# CONFIG
# =========================

MAX_ATTEMPTS = 4
MAX_SYLLABUS_CHARS = 4500
MIN_EXTRACT_CHARS_BEFORE_STOP = 600

STRONG_START_PATTERNS = [
    r"\bunit\s*[ivx\d]+",
    r"\bmodule\s*\d+",
    r"\bchapter\s*\d+",
    r"\bweek\s*\d+",
]

WEAK_START_PATTERNS = [
    r"\bcourse\s+content\b",
    r"\bsyllabus\b"
]

SYLLABUS_STOP_PATTERNS = [
    r"learning\s+outcomes",
    r"course\s+outcomes",
    r"assessment",
    r"evaluation",
    r"references",
    r"textbooks",
    r"recommended\s+books"
]


# =========================
# INPUT
# =========================


def _normalize_line(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip()


def _dedupe_consecutive(lines: list[str]) -> list[str]:
    deduped = []
    prev = None
    for line in lines:
        if not line or line == prev:
            continue
        deduped.append(line)
        prev = line
    return deduped

def read_syllabus(file_path: str) -> str:
    text_lines = []
    
    if file_path.lower().endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        for line in text.splitlines():
                            clean = _normalize_line(line)
                            if clean:
                                text_lines.append(clean)
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {str(e)}")
            
    else:
        # Assume DOCX
        try:
            doc = Document(file_path)

            # Normal paragraphs
            for p in doc.paragraphs:
                clean = _normalize_line(p.text)
                if clean:
                    text_lines.append(clean)

            # Tables (IMPORTANT)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        clean = _normalize_line(cell.text)
                        if clean:
                            text_lines.append(clean)
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {str(e)}")

    return "\n".join(_dedupe_consecutive(text_lines))


# =========================
# SYLLABUS WINDOW EXTRACTION
# =========================

def extract_syllabus_window(raw_text: str) -> str:
    """
    Deterministically finds the syllabus section and
    extracts a bounded window to avoid noise.
    """
    lines = [_normalize_line(line) for line in raw_text.splitlines() if _normalize_line(line)]
    start_idx = 0

    # Prefer the first strong "Unit/Module/Chapter/Week" marker.
    strong_start_idx = None
    for i, line in enumerate(lines):
        for pat in STRONG_START_PATTERNS:
            if re.search(pat, line, re.IGNORECASE):
                strong_start_idx = i
                break
        if strong_start_idx is not None:
            break

    if strong_start_idx is not None:
        start_idx = strong_start_idx
    else:
        # Fall back to weaker "Course Content/Syllabus" markers.
        for i, line in enumerate(lines):
            for pat in WEAK_START_PATTERNS:
                if re.search(pat, line, re.IGNORECASE):
                    start_idx = i
                    break
            if start_idx:
                break

    collected = []
    char_count = 0

    for line in lines[start_idx:]:
        # Stop when we hit non-syllabus sections
        for pat in SYLLABUS_STOP_PATTERNS:
            if re.search(pat, line, re.IGNORECASE):
                if char_count >= MIN_EXTRACT_CHARS_BEFORE_STOP:
                    return "\n".join(collected)
                # Ignore very-early stop markers and continue searching.
                line = ""
                break

        if not line:
            continue

        collected.append(line)
        char_count += len(line)

        if char_count >= MAX_SYLLABUS_CHARS:
            break

    # Safe fallback
    return "\n".join(collected) if collected else raw_text[:MAX_SYLLABUS_CHARS]


# =========================
# QUESTION PAPER GENERATION
# =========================

def generate_question_paper(syllabus_window: str, api_key: str, retry_note: str = "", model_id: str = "mistralai/mistral-7b-instruct", difficulty: str="Medium", exam_format: str="End-Semester", past_papers_text: str="") -> str:
    
    # Customize instructions based on difficulty
    difficulty_instructions = ""
    if difficulty.lower() == "hard":
        difficulty_instructions = "- Emphasize application, analysis, and evaluation level questions (Bloom's Taxonomy).\n- Questions must be challenging and require deep conceptual understanding."
    elif difficulty.lower() == "easy":
        difficulty_instructions = "- Emphasize recall, definitional, and basic understanding level questions.\n- Keep the language simple and direct."
    else:
        difficulty_instructions = "- Provide a balanced mix of recall and application based questions."

    # Customize format based on selection
    format_structure = ""
    if exam_format.lower() == "mid-term":
        format_structure = """
SECTION A
- EXACTLY 3 questions
- Each question carries 5 marks
- No sub-choices

SECTION B
- EXACTLY 2 questions
- Each question carries 10 marks
- Each question MUST have exactly ONE internal choice (OR)

SECTION C
- EXACTLY 1 question
- Carries 15 marks
- MUST have exactly ONE internal choice (OR)
"""
    else:  # End-Semester
        format_structure = """
SECTION A
- EXACTLY 6 questions
- Each question carries 4 marks
- No sub-choices
- No extra questions

SECTION B
- EXACTLY 2 questions
- Each question carries 10 marks
- EACH question MUST have exactly ONE internal choice (OR)
- Do NOT add extra questions or extra ORs

SECTION C
- EXACTLY 1 question
- Carries 16 marks
- MUST have exactly ONE internal choice (OR)
- Do NOT add more than one OR
"""

    past_paper_instruction = ""
    if past_papers_text:
        past_paper_instruction = f"""
STRICT REQUIREMENT FOR UNIQUE QUESTIONS:
The following are questions that have been asked in previous exams for this subject.
YOU MUST NOT REPLICATE OR SUBSTANTIALLY PARAPHRASE ANY OF THESE PAST QUESTIONS.
You must generate ENTIRELY NEW questions testing different aspects or framing scenarios differently.

PAST QUESTIONS:
{past_papers_text[:2000]} # Trim to avoid context limits
"""

    prompt = f"""
You are a university exam paper setter.

{retry_note}

The following text is the EXACT syllabus.
You MUST:
- Use ONLY topics present in the syllabus
- NOT introduce new concepts
- NOT rename or reinterpret the subject
- NOT generalize beyond the syllabus text
{difficulty_instructions}
{past_paper_instruction}

SYLLABUS:
{syllabus_window}

Generate a question paper with EXACTLY this structure:

QUESTION PAPER
SUBJECT (use wording from syllabus)
TIME AND MARKS

STRICT PAPER STRUCTURE (MANDATORY):
{format_structure}

INSTRUCTIONS (at the start only)

STRICTLY FORBIDDEN:
- Markdown (**, *, ---, ###)
- Decorative separators
- Emphasis symbols of any kind

"""
    return generate_text(prompt, api_key=api_key, model_id=model_id)


# =========================
# VALIDATION (FORMAT ONLY)
# =========================

def validate_paper_structure(paper: str) -> list[str]:
    violations = []
    lower = paper.lower()

    required_sections = [
        "section a",
        "section b",
        "section c",
        "instructions"
    ]

    for section in required_sections:
        if section not in lower:
            violations.append(f"Missing {section.upper()}")

    return violations


# =========================
# RETRY ORCHESTRATION
# =========================

def generate_with_retries(raw_syllabus: str, api_key: str, model_id: str = "mistralai/mistral-7b-instruct", difficulty: str="Medium", exam_format: str="End-Semester", past_papers_text: str="") -> str:
    syllabus_window = extract_syllabus_window(raw_syllabus)

    retry_note = ""
    for attempt in range(1, MAX_ATTEMPTS + 1):
        paper = generate_question_paper(syllabus_window, api_key, retry_note, model_id, difficulty, exam_format, past_papers_text)
        violations = validate_paper_structure(paper)

        if not violations:
            return paper

        retry_note = (
            "The previous attempt violated the required format:\n- "
            + "\n- ".join(violations)
            + "\nRegenerate while strictly following the structure."
        )

    raise RuntimeError("Failed to generate a valid question paper.")

# =========================
# PAPER ANALYTICS
# =========================
import json

def analyze_paper_quality(syllabus_text: str, paper_text: str, api_key: str, model_id: str = "mistralai/mistral-7b-instruct") -> dict:
    prompt = f"""
You are an expert academic evaluator. Your job is to analyze the generated exam paper against its provided syllabus and output ONLY a valid JSON object analyzing its quality.

SYLLABUS:
{syllabus_text}

EXAM PAPER:
{paper_text}

Analyze the exam paper based on:
1. Topic Coverage: What percentage (0-100) of the syllabus topics are fairly represented?
2. Difficulty Score: Rate the overall difficulty from 1 (Very Easy) to 10 (Extremely Hard).
3. Red Flags: Are there any questions that clearly fall OUTSIDE the scope of the syllabus?

Output EXACTLY AND ONLY this JSON format:
{{
    "coverage_percentage": 85,
    "difficulty_score": 7,
    "red_flags": ["Question 3a is about quantum mechanics, which is not in the syllabus.", "None otherwise."]
}}
"""
    try:
        response_text = generate_text(prompt, api_key=api_key, model_id=model_id)
        # Strip potential markdown blocks if the LLM adds them
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]
            
        return json.loads(response_text)
    except Exception as e:
        print(f"Analytics Error: {e}")
        return {
            "coverage_percentage": "N/A",
            "difficulty_score": "N/A",
            "red_flags": ["Failed to analyze paper. LLM response was not valid JSON."]
        }

# =========================
# OUTPUT
# =========================

def save_to_docx(text: str, output_path: str, subject_name: str = "Subject Name", exam_format: str = "End-Semester"):
    doc = Document()

    # 1. UNIVERSITY HEADER
    heading = doc.add_paragraph()
    heading.alignment = 1 # Center
    
    run1 = heading.add_run("GLOBAL UNIVERSITY OF TECHNOLOGY\n")
    run1.bold = True
    run1.font.size = 204800 # ~16pt
    
    run2 = heading.add_run(f"{exam_format.upper()} EXAMINATION\n")
    run2.bold = True
    
    run3 = heading.add_run(f"Course Title: {subject_name}\n")
    
    doc.add_paragraph("-" * 60).alignment = 1
    
    # Details Row
    details = doc.add_paragraph()
    details.add_run(f"Time: {'3 Hours' if exam_format == 'End-Semester' else '2 Hours'}").bold = True
    details.add_run("\t\t\t\t\t\t") # basic tabs for spacing
    details.add_run(f"Max Marks: {'100' if exam_format == 'End-Semester' else '50'}").bold = True
    
    doc.add_paragraph("-" * 60).alignment = 1

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.upper() in [
            "QUESTION PAPER",
            "SUBJECT",
            "TIME AND MARKS",
            "SECTION A",
            "SECTION B",
            "SECTION C",
            "INSTRUCTIONS"
        ]:
            # Skip the ones we hardcoded in the header
            if line.upper() in ["QUESTION PAPER", "SUBJECT", "TIME AND MARKS"]:
                continue
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
